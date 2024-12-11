# import streamlit as st
# import datetime
# import sqlite3
# import requests
# from goose3 import Goose
# from io import BytesIO
# from docx import Document
# import urllib.parse
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# import docx

# # Title for the monitoring page
# st.title("Real-Time Monitoring and Data Collection")

# # Database connection
# def get_connection():
#     return sqlite3.connect("reputation_management.db")

# # Google Custom Search API function
# def google_search(query, date_filter_type=None, date_filter_value=None):
#     API_KEY = 'AIzaSyC_0LAqVIA0Z7YLbbWKSOHxY0_sMaqQAko'
#     CSE_ID = 'b13bba6a528214af0'
#     search_url = "https://www.googleapis.com/customsearch/v1"
    
#     dateRestrict = ""
#     if date_filter_type == "Days":
#         dateRestrict = f"d{date_filter_value}"
#     elif date_filter_type == "Weeks":
#         dateRestrict = f"w{date_filter_value}"
#     elif date_filter_type == "Months":
#         dateRestrict = f"m{date_filter_value}"
#     elif date_filter_type == "Years":
#         dateRestrict = f"y{date_filter_value}"
    
#     params = {
#         "key": API_KEY,
#         "cx": CSE_ID,
#         "q": f'allintitle:{query}',
#         "num": 5,
#         "dateRestrict": dateRestrict
#     }
    
#     response = requests.get(search_url, params=params)
#     return response.json() if response.status_code == 200 else None

# # Extract full content
# def extract_full_content(url):
#     try:
#         g = Goose()
#         article = g.extract(url=url)
#         return article.cleaned_text
#     except:
#         return "Could not extract content."

# # Function to clean the links
# def clean_link(link):
#     try:
#         cleaned = urllib.parse.unquote(link, encoding='utf-8')
#     except UnicodeDecodeError:
#         cleaned = urllib.parse.unquote(link, encoding='iso-8859-1')
#     return cleaned

# # Save monitoring data in the database
# def save_monitoring_data(results, query):
#     conn = get_connection()
#     cursor = conn.cursor()
#     for result in results:
#         if not result["content"] == "" or "Could not extract content.":
#             cursor.execute('''
#                 INSERT INTO monitoring_data (source, content, timestamp, query)
#                 VALUES (?, ?, ?, ?)
#             ''', ("Google Search", result["content"], result["timestamp"], query))
#     conn.commit()
#     conn.close()

# # Function to add hyperlink in DOCX
# def add_hyperlink(paragraph, url, text):
#     part = paragraph.part
#     r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
#     hyperlink = OxmlElement('w:hyperlink')
#     hyperlink.set(qn('r:id'), r_id)
#     new_run = OxmlElement('w:r')
#     r_pr = OxmlElement('w:rPr')
#     new_run.append(r_pr)
#     new_run.text = text
#     hyperlink.append(new_run)
#     paragraph._element.append(hyperlink)
#     return paragraph

# # Function to create a DOCX file with monitoring data
# def create_docx_file(results):
#     doc = Document()
#     doc.add_heading("Monitoring Results", 0)
    
#     for result in results:
#         doc.add_heading(result["title"], level=1)
#         p = doc.add_paragraph("Link: ")
#         add_hyperlink(p, result["link"], result["link"])
#         doc.add_paragraph(f"Snippet: {result['snippet']}")
#         doc.add_paragraph("Content:")
#         doc.add_paragraph(result["content"])
#         doc.add_page_break()
    
#     buffer = BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer

# # Input fields for monitoring criteria
# search_query = st.text_input("Enter keywords for tracking")
# date_filter_type = st.selectbox("Choose Time Range", ["No Filter", "Days", "Weeks", "Months", "Years"])
# date_filter_value = st.number_input("Specify the number of selected units (e.g., days, weeks)", min_value=1, value=1)
# start_monitoring = st.button("Start Monitoring")

# # Fetch results if "Start Monitoring" is clicked
# if start_monitoring and search_query:
#     st.write(f"Monitoring for **{search_query}**...")
#     results = google_search(search_query, date_filter_type, date_filter_value)
    
#     if results:
#         # Store results in session state
#         st.session_state["results_data"] = [
#             {
#                 "title": item['title'],
#                 "link": clean_link(item['link']),
#                 "snippet": item['snippet'],
#                 "content": extract_full_content(item['link']),
#                 "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#             } for item in results.get("items", [])
#         ]
#         st.success("Results fetched successfully.")
#     else:
#         st.warning("No results found.")

# # Display results if available
# if "results_data" in st.session_state:
#     for item in st.session_state["results_data"]:
#         st.write(f"**Title:** {item['title']}")
#         st.write(f"**Link:** {item['link']}")
#         st.write(f"**Snippet:** {item['snippet']}")
#         st.write(f"**Content:** {item['content']}")
#         st.write("---")

# # Button to save all results in the database
# if st.button("Save All Results to Database"):
#     if "results_data" in st.session_state:
#         save_monitoring_data(st.session_state["results_data"], search_query)
#         st.success("All results saved to the database.")
#     else:
#         st.warning("No results to save.")

# # Button to create a DOCX file for all results
# if st.button("Download All Results as DOCX"):
#     if "results_data" in st.session_state:
#         docx_content = create_docx_file(st.session_state["results_data"])
#         st.download_button(
#             label="Download Monitoring Results as DOCX",
#             data=docx_content,
#             file_name="monitoring_results.docx",
#             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         )
#     else:
#         st.warning("No results to download.")


# import streamlit as st
# import datetime
# import sqlite3
# import requests
# from goose3 import Goose
# from io import BytesIO
# from docx import Document
# import urllib.parse
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# import docx

# # Page Configuration
# # st.set_page_config(
# #     page_title="Real-Time Monitoring",
# #     layout="wide"
# # )

# # Database Connection
# def get_connection():
#     return sqlite3.connect("reputation_management.db")

# # Google Custom Search API Function
# def google_search(query, date_filter_type=None, date_filter_value=None):
#     API_KEY = 'AIzaSyC_0LAqVIA0Z7YLbbWKSOHxY0_sMaqQAko'
#     CSE_ID = 'b13bba6a528214af0'
#     search_url = "https://www.googleapis.com/customsearch/v1"
    
#     dateRestrict = ""
#     if date_filter_type == "Days":
#         dateRestrict = f"d{date_filter_value}"
#     elif date_filter_type == "Weeks":
#         dateRestrict = f"w{date_filter_value}"
#     elif date_filter_type == "Months":
#         dateRestrict = f"m{date_filter_value}"
#     elif date_filter_type == "Years":
#         dateRestrict = f"y{date_filter_value}"
    
#     params = {
#         "key": API_KEY,
#         "cx": CSE_ID,
#         "q": f'allintitle:{query}',
#         "num": 5,
#         "dateRestrict": dateRestrict
#     }
    
#     response = requests.get(search_url, params=params)
#     return response.json() if response.status_code == 200 else None

# # Extract Full Content Function
# def extract_full_content(url):
#     try:
#         g = Goose()
#         article = g.extract(url=url)
#         return article.cleaned_text
#     except:
#         return "Could not extract content."

# # Clean Link Function
# def clean_link(link):
#     try:
#         cleaned = urllib.parse.unquote(link, encoding='utf-8')
#     except UnicodeDecodeError:
#         cleaned = urllib.parse.unquote(link, encoding='iso-8859-1')
#     return cleaned

# # Save Monitoring Data in the Database
# def save_monitoring_data(results, query):
#     conn = get_connection()
#     cursor = conn.cursor()
#     for result in results:
#         if not result["content"] == "" or "Could not extract content.":
#             cursor.execute('''
#                 INSERT INTO monitoring_data (source, content, timestamp, query)
#                 VALUES (?, ?, ?, ?)
#             ''', ("Google Search", result["content"], result["timestamp"], query))
#     conn.commit()
#     conn.close()

# # Add Hyperlink to DOCX
# def add_hyperlink(paragraph, url, text):
#     part = paragraph.part
#     r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
#     hyperlink = OxmlElement('w:hyperlink')
#     hyperlink.set(qn('r:id'), r_id)
#     new_run = OxmlElement('w:r')
#     r_pr = OxmlElement('w:rPr')
#     new_run.append(r_pr)
#     new_run.text = text
#     hyperlink.append(new_run)
#     paragraph._element.append(hyperlink)
#     return paragraph

# # Create DOCX File Function
# def create_docx_file(results):
#     doc = Document()
#     doc.add_heading("Monitoring Results", 0)
    
#     for result in results:
#         doc.add_heading(result["title"], level=1)
#         p = doc.add_paragraph("Link: ")
#         add_hyperlink(p, result["link"], result["link"])
#         doc.add_paragraph(f"Snippet: {result['snippet']}")
#         doc.add_paragraph("Content:")
#         doc.add_paragraph(result["content"])
#         doc.add_page_break()
    
#     buffer = BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer

# # UI Layout with Tabs
# st.title("Real-Time Monitoring and Data Collection")
# tabs = st.tabs(["Live Monitoring", "Custom Search", "Saved Results"])

# # Tab 1: Live Monitoring
# with tabs[0]:
#     st.subheader("Live Monitoring Feed")
#     st.markdown("Real-time updates from configured sources will be displayed here.")
#     feed_placeholder = st.empty()

# # Tab 2: Custom Search
# with tabs[1]:
#     st.subheader("Perform Custom Search")
#     search_query = st.text_input("Enter keywords for tracking")
#     date_filter_type = st.selectbox("Choose Time Range", ["No Filter", "Days", "Weeks", "Months", "Years"])
#     date_filter_value = st.number_input("Specify the number of selected units", min_value=1, value=1)
#     start_monitoring = st.button("Start Monitoring")
    
#     if start_monitoring and search_query:
#         st.write(f"Monitoring for **{search_query}**...")
#         results = google_search(search_query, date_filter_type, date_filter_value)
        
#         if results:
#             st.session_state["results_data"] = [
#                 {
#                     "title": item['title'],
#                     "link": clean_link(item['link']),
#                     "snippet": item['snippet'],
#                     "content": extract_full_content(item['link']),
#                     "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#                 } for item in results.get("items", [])
#             ]
#             st.success("Results fetched successfully.")
#         else:
#             st.warning("No results found.")
    
#     # Display Results
#     if "results_data" in st.session_state:
#         for item in st.session_state["results_data"]:
#             st.write(f"**Title:** {item['title']}")
#             st.write(f"**Link:** {item['link']}")
#             st.write(f"**Snippet:** {item['snippet']}")
#             st.write(f"**Content:** {item['content']}")
#             st.write("---")

# # Tab 3: Saved Results
# with tabs[2]:
#     if "results_data" in st.session_state:
#         if st.button("Save All Results to Database"):
#             save_monitoring_data(st.session_state["results_data"], search_query)
#             st.success("All results saved to the database.")
        
#         if st.button("Download All Results as DOCX"):
#             docx_content = create_docx_file(st.session_state["results_data"])
#             st.download_button(
#                 label="Download Monitoring Results as DOCX",
#                 data=docx_content,
#                 file_name="monitoring_results.docx",
#                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#             )
#     else:
#         st.warning("No results to display.")

import streamlit as st
import datetime
import sqlite3
import requests
from goose3 import Goose
from io import BytesIO
from docx import Document
import urllib.parse
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx
import time
from streamlit_option_menu import option_menu

# Page Configuration
# st.set_page_config(
#     page_title="Real-Time Monitoring",
#     layout="wide",
#     page_icon="🌍"
# )

# Custom CSS Styles
st.markdown("""
<style>
    .main {
        background-color: #f7f8fc;
    }
    .card {
        background-color: white;
        border-radius: 10px;
        box-shadow: 0px 4px 10px rgba(0, 0, 0, 0.1);
        padding: 20px;
        margin-bottom: 20px;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .card:hover {
        transform: translateY(-5px);
        box-shadow: 0px 6px 15px rgba(0, 0, 0, 0.15);
    }
    .card-header {
        font-size: 1.5rem;
        font-weight: bold;
        color: #007BFF;
    }
    .card-link {
        color: #0056b3;
        text-decoration: none;
    }
    .card-link:hover {
        text-decoration: underline;
    }
    .card-content {
        margin-top: 10px;
        font-size: 1rem;
        color: #555;
    }
    .card-timestamp {
        font-size: 0.9rem;
        color: #999;
        margin-top: 15px;
    }
    div.stButton > button {
        background-color: #007BFF;
        color: white;
        border: none;
        padding: 10px 20px;
        border-radius: 5px;
        cursor: pointer;
        font-size: 16px;
    }
    div.stButton > button:hover {
        background-color: #0056b3;
    }
</style>
""", unsafe_allow_html=True)

# Database Connection
def get_connection():
    return sqlite3.connect("reputation_management.db")

# # Google Custom Search API Function
# def google_search(query, date_filter_type=None, date_filter_value=None):
#     API_KEY = 'AIzaSyC_0LAqVIA0Z7YLbbWKSOHxY0_sMaqQAko'
#     CSE_ID = 'b13bba6a528214af0'
#     search_url = "https://www.googleapis.com/customsearch/v1"
    
#     dateRestrict = ""
#     if date_filter_type == "Days":
#         dateRestrict = f"d{date_filter_value}"
#     elif date_filter_type == "Weeks":
#         dateRestrict = f"w{date_filter_value}"
#     elif date_filter_type == "Months":
#         dateRestrict = f"m{date_filter_value}"
#     elif date_filter_type == "Years":
#         dateRestrict = f"y{date_filter_value}"
    
#     params = {
#         "key": API_KEY,
#         "cx": CSE_ID,
#         "q": f'allintitle:{query}',
#         "num": 5,
#         "dateRestrict": dateRestrict
#     }
    
#     response = requests.get(search_url, params=params)
#     return response.json() if response.status_code == 200 else None

# # Extract Full Content Function
# def extract_full_content(url):
#     try:
#         g = Goose()
#         article = g.extract(url=url)
#         return article.cleaned_text
#     except:
#         return "Could not extract content."

# Google Custom Search API Function
def google_search(query, date_filter_type=None, date_filter_value=None):
    API_KEY = 'AIzaSyC_0LAqVIA0Z7YLbbWKSOHxY0_sMaqQAko'
    CSE_ID = 'b13bba6a528214af0'
    search_url = "https://www.googleapis.com/customsearch/v1"
    
    dateRestrict = ""
    if date_filter_type == "Days":
        dateRestrict = f"d{date_filter_value}"
    elif date_filter_type == "Weeks":
        dateRestrict = f"w{date_filter_value}"
    elif date_filter_type == "Months":
        dateRestrict = f"m{date_filter_value}"
    elif date_filter_type == "Years":
        dateRestrict = f"y{date_filter_value}"
    
    params = {
        "key": API_KEY,
        "cx": CSE_ID,
        "q": query,
        "num": 10,
        "dateRestrict": dateRestrict
    }
    
    response = requests.get(search_url, params=params)
    return response.json() if response.status_code == 200 else None

# Extract Full Content Function
def extract_full_content(url):
    try:
        g = Goose()
        article = g.extract(url=url)
        return article.cleaned_text
    except:
        return "Could not extract content."

# Clean Link Function
def clean_link(link):
    try:
        cleaned = urllib.parse.unquote(link, encoding='utf-8')
    except UnicodeDecodeError:
        cleaned = urllib.parse.unquote(link, encoding='iso-8859-1')
    return cleaned

# Save Monitoring Data in the Database
def save_monitoring_data(results, query):
    conn = get_connection()
    cursor = conn.cursor()
    for result in results:
        if not result["content"] == "" or "Could not extract content.":
            cursor.execute('''
                INSERT INTO monitoring_data (source, content, timestamp, query)
                VALUES (?, ?, ?, ?)
            ''', ("Google Search", result["content"], result["timestamp"], query))
    conn.commit()
    conn.close()

# Add Hyperlink to DOCX
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return paragraph

# Create DOCX File Function
def create_docx_file(results):
    doc = Document()
    doc.add_heading("Monitoring Results", 0)
    
    for result in results:
        doc.add_heading(result["title"], level=1)
        p = doc.add_paragraph("Link: ")
        add_hyperlink(p, result["link"], result["link"])
        doc.add_paragraph(f"Snippet: {result['snippet']}")
        doc.add_paragraph("Content:")
        doc.add_paragraph(result["content"])
        doc.add_page_break()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Main Content with Tabs
st.title("🌍 Real-Time Monitoring and Data Collection")

# tabs = st.tabs(["Monitoring", "Live Feed"])
selected_tab = option_menu(
    menu_title=None,
    options=["Monitoring", "Live Feed"],
    icons=["monitoring", "live-feed"],
    default_index=0,
    orientation="horizontal",
)
if selected_tab == "Monitoring":
    # Monitoring Tab
    search_query = st.text_input("🔍 Enter Keywords for Monitoring")
    date_filter_type = st.selectbox("📅 Choose Time Range", ["No Filter", "Days", "Weeks", "Months", "Years"])
    date_filter_value = st.number_input("Specify Units", min_value=1, value=1)
    start_monitoring = st.button("Start Monitoring 🚀", key="monitoring_button")

    if start_monitoring and search_query:
        st.write(f"Searching for **{search_query}**...")
        results = google_search(search_query, date_filter_type, date_filter_value)
        print(results)

        if results:
            st.session_state["results_data"] = [
                {
                    "title": item['title'],
                    "link": clean_link(item['link']),
                    "snippet": item['snippet'],
                    "content": extract_full_content(item['link']),
                    "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                } for item in results.get("items", [])
            ]
            st.success("✅ Results fetched successfully!")
        else:
            st.warning("⚠️ No results found.")

    # Display Results in Cards
    if "results_data" in st.session_state:
        for item in st.session_state["results_data"]:
            st.markdown(f"""
            <div class="card">
                <div class="card-header">{item['title']}</div>
                <p><strong>Link:</strong> <a class="card-link" href="{item['link']}" target="_blank">{item['link']}</a></p>
                <p class="card-content"><strong>Snippet:</strong> {item['snippet']}</p>
                <p class="card-content"><strong>Content:</strong> {item['content'][:200]}...</p>
                <p class="card-timestamp">Timestamp: {item['timestamp']}</p>
            </div>
            """, unsafe_allow_html=True)

    # Save and Download Options
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Save Results to Database", key="save_button"):
            save_monitoring_data(st.session_state["results_data"], search_query)
            st.success("✅ Results saved successfully!")

    with col2:
        if st.button("📄 Download Results as DOCX", key="download_button"):
            docx_content = create_docx_file(st.session_state["results_data"])
            st.download_button(
                label="Download Monitoring Results 📥",
                data=docx_content,
                file_name="monitoring_results.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

elif selected_tab == "Live Feed":
    # Live Feed Tab
    st.header("📡 Live Feed")
    st.write("This section will display a real-time feed of monitored data or updates.")
    # Placeholder for real-time feed content
    feed_placeholder = st.empty()
    # Example: Displaying a live timestamp
    while True:
        current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        feed_placeholder.write(f"Live Feed Update: {current_time}")
        time.sleep(1)  # Simulate real-time updates
