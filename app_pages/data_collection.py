import streamlit as st
import requests
import datetime
import sqlite3
from bs4 import BeautifulSoup
from goose3 import Goose
from app_pages.db_setup import create_db
from io import BytesIO
from docx import Document
import urllib.parse

# Title and page setup
st.title("Company Reputation Monitoring")
st.header("Data Collection and Analysis")

# Input fields for search criteria
company_name = st.text_input("Company Name", "")
search_group = st.text_input("Search Group", "")
sources = st.text_area("Enter news sources (comma-separated)", "")
classification = st.selectbox("Classify sources", ["Public", "Competitors", "Allies", "Enemies"])
date_filter_type = st.selectbox("Choose Time Range", ["No Filter", "Days", "Weeks", "Months", "Years"])
date_filter_value = st.number_input("Specify the number of selected units (e.g., days, weeks)", min_value=1, value=1)
submit_button = st.button("Search")
save_all_button = st.button("Save All Results")  # New Save All button
create_db()

# Google Custom Search API function
def google_search(query, site=None, date_filter_type=None, date_filter_value=None):
    API_KEY = 'AIzaSyC_0LAqVIA0Z7YLbbWKSOHxY0_sMaqQAko'
    CSE_ID = 'b13bba6a528214af0'
    search_url = "https://www.googleapis.com/customsearch/v1"
    
    dateRestrict = ""
    if date_filter_type == "Days":
        dateRestrict = f"d{date_filter_value}"
    elif date_filter_type == "Weeks":
        dateRestrict = f"w{date_filter_value}"
    elif date_filter_type == "Months":
        dateRestrict = f"m{date_filter_value}"
    elif date_filter_type == "Years":
        dateRestrict = f"y{date_filter_value}"
    
    params = {
        "key": API_KEY,
        "cx": CSE_ID,
        "q": f'allintitle:{query}',
        "num": 5,
        "dateRestrict": dateRestrict,
        "siteSearch": site if site else None,
    }
    
    response = requests.get(search_url, params=params)
    return response.json() if response.status_code == 200 else None

# Extract full content
def extract_full_content(url):
    try:
        g = Goose()
        article = g.extract(url=url) 
        return article.cleaned_text
    except:
        return "Could not extract content."

# Function to clean the links
def clean_link(link):
    try:
            # Decode the URL using utf-8 encoding (common for web pages)
            cleaned = urllib.parse.unquote(link, encoding='utf-8')
    except UnicodeDecodeError:
            # If decoding with utf-8 fails, try using iso-8859-1 (another common encoding)
            cleaned = urllib.parse.unquote(link, encoding='iso-8859-1')
    return cleaned

# Save search results in the database
def save_search_to_db(company_name, source, title, link, snippet, full_content, classification, date, search_group):
    conn = sqlite3.connect('company_reputation.db')
    c = conn.cursor()
    c.execute('''INSERT INTO searches (company_name, source, title, link, snippet, full_content, classification, date, search_group)
             VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''', 
          (company_name, source, title, link, snippet, full_content, classification, date, search_group))
    conn.commit()
    conn.close()

# Initialize session state for all_results
if 'all_results' not in st.session_state:
    st.session_state.all_results = []

# Display search results and save them
if submit_button:
    st.write(f"Searching for news on **{company_name}**...")
    st.session_state.all_results.clear()  # Clear previous results for each new search
    
    source_list = [source.strip() for source in sources.split(",") if source.strip()]
    for source in source_list:
        results = google_search(company_name, site=source, date_filter_type=date_filter_type, date_filter_value=date_filter_value)
        if results:
            for item in results.get("items", []):
                title = item['title']
                link = clean_link(item['link'])
                snippet = item['snippet']
                date = item.get('pagemap', {}).get('metatags', [{}])[0].get('publisheddate')
                full_content = extract_full_content(link)
                
                # Append each result to session state for persistence
                st.session_state.all_results.append((company_name, source, title, link, snippet, full_content, classification, date, search_group))

                # Display the result
                st.write(f"**Title:** {title}")
                st.write(f"**Link:** {link}")
                st.write(f"**Snippet:** {snippet}")
                st.write("---")

# Save all results if the user clicks "Save All"
if save_all_button:
    if st.session_state.all_results:
        for result in st.session_state.all_results:
            save_search_to_db(*result)
        st.success("All results have been saved.")
    else:
        st.warning("No results to save.")

# Function to create a TXT file
def create_txt_file(results):
    content = ""
    for result in results:
        content += f"Title: {result[2]}\nLink: {result[3]}\nSnippet: {result[4]}\nFull Content: {result[5]}\n\n---\n\n"
    return content

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx
# Function to create a hyperlink in a docx document
def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.
    :param paragraph: The docx paragraph object where the link should be added.
    :param url: The web link (string) to add.
    :param text: The text that will display for the link.
    :return: The modified paragraph object with a hyperlink.
    """
    # Create the hyperlink tag
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    # Join all the parts together
    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)

    # Append the hyperlink to the paragraph
    paragraph._element.append(hyperlink)

    return paragraph

# Function to create a DOCX file with clickable links
def create_docx_file(results):
    doc = Document()
    doc.add_heading("Search Results", 0)

    for result in results:
        doc.add_heading(result[2], level=1)  # Add title
        p = doc.add_paragraph("Link: ")
        add_hyperlink(p, result[3], result[3])  # Add clickable link
        doc.add_paragraph(f"Snippet: {result[4]}")  # Add snippet
        doc.add_paragraph("Full Content:")  # Add full content label
        doc.add_paragraph(result[5])  # Add the full content
        doc.add_page_break()

    # Save to in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Download buttons for TXT and DOCX
if st.session_state.all_results:
    txt_content = create_txt_file(st.session_state.all_results)
    st.download_button(
        label="Download Results as TXT",
        data=txt_content,
        file_name="search_results.txt",
        mime="text/plain"
    )
    
    docx_content = create_docx_file(st.session_state.all_results)
    st.download_button(
        label="Download Results as DOCX",
        data=docx_content,
        file_name="search_results.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )